#!/usr/bin/env python3
"""Perform structural DOCX checks without printing document text."""

from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Check DOCX structure, media, and forbidden terms.")
    parser.add_argument("docx", help="Path to the DOCX file.")
    parser.add_argument("--expected-images", type=int, help="Require this exact number of embedded images.")
    parser.add_argument("--forbidden-term", action="append", default=[], help="Fail if this exact text occurs in the document.")
    parser.add_argument("--report", help="Optional JSON report output path.")
    return parser.parse_args()


def paragraph_texts(document: Document) -> Iterable[str]:
    for paragraph in document.paragraphs:
        yield paragraph.text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text


def main() -> int:
    args = parse_args()
    path = Path(args.docx).resolve()
    report: dict[str, object] = {
        "path": str(path),
        "openable": False,
        "embedded_images": 0,
        "paragraphs": 0,
        "tables": 0,
        "forbidden_matches": [],
        "passed": False,
    }
    errors: list[str] = []
    if not path.is_file():
        errors.append("DOCX file does not exist.")
    elif not zipfile.is_zipfile(path):
        errors.append("File is not a valid OOXML ZIP package.")
    else:
        try:
            with zipfile.ZipFile(path) as archive:
                names = archive.namelist()
                if "word/document.xml" not in names:
                    errors.append("DOCX package lacks word/document.xml.")
                report["embedded_images"] = sum(
                    1 for name in names if name.startswith("word/media/") and not name.endswith("/")
                )
            document = Document(path)
            report["openable"] = True
            report["paragraphs"] = len(document.paragraphs)
            report["tables"] = len(document.tables)
            joined_text = "\n".join(paragraph_texts(document))
            matches = [term for term in args.forbidden_term if term and term in joined_text]
            report["forbidden_matches"] = matches
            if matches:
                errors.append("Forbidden terms found.")
        except Exception as error:  # python-docx and zipfile surface varied exceptions.
            errors.append(f"DOCX reopen failed: {type(error).__name__}.")
    if args.expected_images is not None:
        if args.expected_images < 0:
            errors.append("--expected-images must not be negative.")
        elif report["embedded_images"] != args.expected_images:
            errors.append("Embedded image count differs from --expected-images.")
    report["passed"] = not errors
    report["errors"] = errors
    output = json.dumps(report, ensure_ascii=False, indent=2)
    if args.report:
        report_path = Path(args.report).resolve()
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(output + "\n", encoding="utf-8")
    print(output)
    return 0 if report["passed"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
