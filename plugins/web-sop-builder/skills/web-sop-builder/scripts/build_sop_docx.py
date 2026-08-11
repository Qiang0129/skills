#!/usr/bin/env python3
"""Build a Chinese web-operation SOP DOCX from sop-job.json."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Microsoft YaHei"
INK = "27364A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "6B778C"
HEADER_FILL = "E8EEF5"
NOTE_FILL = "F4F6F9"
CAUTION_FILL = "FFF4D6"
WHITE = "FFFFFF"
GRID = "D5DDE7"
SECRET_ASSIGNMENT = re.compile(
    r"(?ix)(?:password|passwd|pwd|token|api[ _-]?key|secret|credential|密码|口令|令牌|密钥)\s*[:：=]\s*\S+"
)
PROHIBITED_FIELD = re.compile(
    r"(?ix)^(?:password|passwd|pwd|token|api[ _-]?key|secret|credential|密码|口令|令牌|密钥|"
    r"(?:internal[ _-]?)?(?:config|configuration|warehouse|customer|tenant|organization)[ _-]?(?:id|code|value)|"
    r"(?:内部)?配置(?:id|编号|值)|(?:仓库|客户|租户|组织)(?:id|编号|代码))$"
)


class JobError(ValueError):
    """Raise a concise validation error without echoing job contents."""


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a Chinese web-operation SOP DOCX.")
    parser.add_argument("job", nargs="?", default="sop-job.json", help="Path to sop-job.json (default: sop-job.json).")
    return parser.parse_args()


def load_job(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as error:
        raise JobError(f"Job file does not exist: {path}") from error
    except json.JSONDecodeError as error:
        raise JobError(f"Invalid JSON in {path}: line {error.lineno}, column {error.colno}") from error
    if not isinstance(data, dict):
        raise JobError("The job root must be a JSON object.")
    return data


def text(value: Any, field: str, required: bool = True) -> str:
    if value is None and not required:
        return ""
    if not isinstance(value, str) or not value.strip():
        raise JobError(f"{field} must be a non-empty string.")
    return value.strip()


def text_array(value: Any, field: str, required: bool = False) -> list[str]:
    if value is None and not required:
        return []
    if not isinstance(value, list):
        raise JobError(f"{field} must be an array.")
    return [text(item, f"{field}[{index}]") for index, item in enumerate(value)]


def resolve_path(job_path: Path, value: Any, field: str) -> Path:
    raw = text(value, field)
    candidate = Path(raw)
    return candidate if candidate.is_absolute() else job_path.parent / candidate


def reject_credentials(job: dict[str, Any]) -> None:
    def inspect(value: Any) -> None:
        if isinstance(value, str):
            if SECRET_ASSIGNMENT.search(value):
                raise JobError("The job appears to include a credential assignment. Remove credentials before building the SOP.")
        elif isinstance(value, dict):
            for key, item in value.items():
                if PROHIBITED_FIELD.fullmatch(str(key).strip()):
                    raise JobError("The job contains a prohibited credential or internal-configuration field.")
                inspect(item)
        elif isinstance(value, list):
            for item in value:
                inspect(item)

    inspect(job)


def set_run_font(run: Any, size: float = 11, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(paragraph: Any, before: float = 0, after: float = 6, line: float = 1.25, align: Any = None) -> None:
    format_ = paragraph.paragraph_format
    format_.space_before = Pt(before)
    format_.space_after = Pt(after)
    format_.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_run(paragraph: Any, value: str, size: float = 11, color: str = INK, bold: bool = False, italic: bool = False) -> Any:
    run = paragraph.add_run(value)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_border(cell: Any, color: str = GRID, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prepare_table(table: Any, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_width = tc_pr.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                tc_pr.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header(row: Any) -> None:
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(node)


def clear_cell(cell: Any) -> Any:
    cell.text = ""
    return cell.paragraphs[0]


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    for style_name, size, color, bold in (("Normal", 10.5, INK, False), ("Caption", 9, MUTED, False)):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
    for level, size in ((1, 16), (2, 13), (3, 11.5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)
        style.font.bold = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(header, after=0, line=1.0)
    add_run(header, title, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(footer, after=0, line=1.0)
    add_run(footer, "内部操作指南", size=8.5, color=MUTED)


def add_heading(doc: Document, value: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph(paragraph, before={1: 18, 2: 13, 3: 9}[level], after={1: 8, 2: 6, 3: 4}[level], line=1.2)
    add_run(paragraph, value, size={1: 16, 2: 13, 3: 11.5}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)


def add_body(doc: Document, value: str, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=after)
    add_run(paragraph, value)


def add_bullet(doc: Document, value: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph(paragraph, after=4)
    add_run(paragraph, value)


def add_note(doc: Document, label: str, value: str, fill: str = NOTE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    prepare_table(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    paragraph = clear_cell(cell)
    set_paragraph(paragraph, after=0, line=1.2)
    add_run(paragraph, f"{label}  ", size=10, color=DARK_BLUE, bold=True)
    add_run(paragraph, value, size=10)
    doc.add_paragraph()


def add_figure(doc: Document, job_path: Path, figure: dict[str, Any]) -> None:
    figure_path = resolve_path(job_path, figure.get("path"), "figure.path")
    if not figure_path.is_file():
        raise JobError(f"Reviewed figure does not exist: {figure_path}")
    caption = text(figure.get("caption"), "figure.caption")
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=2, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run()
    run.add_picture(str(figure_path), width=Inches(6.4))
    doc_pr = run._element.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", caption)
    caption_paragraph = doc.add_paragraph(style="Caption")
    set_paragraph(caption_paragraph, before=1, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(caption_paragraph, caption, size=9, color=MUTED, italic=True)


def add_cover(doc: Document, job: dict[str, Any]) -> None:
    subtitle = text(job.get("subtitle"), "subtitle", required=False)
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=70, after=12, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(paragraph, text(job["system_name"], "system_name"), size=11, color=DARK_BLUE, bold=True)
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(paragraph, text(job["title"], "title"), size=26, color=INK, bold=True)
    if subtitle:
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, after=25, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(paragraph, subtitle, size=14, color=BLUE, bold=True)
    rows = [
        ("适用系统", text(job["system_name"], "system_name")),
        ("适用范围", text(job["applicable_scope"], "applicable_scope")),
        ("适用人员", text(job["applicable_users"], "applicable_users")),
        ("交付说明", "基于已确认的页面操作和标注截图编写；不包含登录凭据或内部配置值。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    prepare_table(table, [2200, 7160])
    for row, (label, value) in zip(table.rows, rows):
        shade(row.cells[0], HEADER_FILL)
        for cell, item, style in ((row.cells[0], label, True), (row.cells[1], value, False)):
            paragraph = clear_cell(cell)
            set_paragraph(paragraph, after=0, line=1.15)
            add_run(paragraph, item, size=9.6, color=DARK_BLUE if style else INK, bold=style)
    add_note(doc, "使用说明", "本文件用于内部操作培训与复核。外发或生产环境材料应先确认数据处理要求。", CAUTION_FILL)
    doc.add_page_break()


def validate_sections(job: dict[str, Any]) -> list[dict[str, Any]]:
    sections = job.get("sections")
    if not isinstance(sections, list) or not sections:
        raise JobError("sections must contain at least one section.")
    for section_index, section in enumerate(sections):
        if not isinstance(section, dict):
            raise JobError(f"sections[{section_index}] must be an object.")
        text(section.get("title"), f"sections[{section_index}].title")
        steps = section.get("steps")
        if not isinstance(steps, list) or not steps:
            raise JobError(f"sections[{section_index}].steps must contain at least one step.")
        for step_index, step in enumerate(steps):
            if not isinstance(step, dict):
                raise JobError(f"sections[{section_index}].steps[{step_index}] must be an object.")
            prefix = f"sections[{section_index}].steps[{step_index}]"
            for field in ("title", "action", "expected_result"):
                text(step.get(field), f"{prefix}.{field}")
            text_array(step.get("notes"), f"{prefix}.notes")
            figures = step.get("figures", [])
            if not isinstance(figures, list):
                raise JobError(f"{prefix}.figures must be an array.")
            for figure_index, figure in enumerate(figures):
                if not isinstance(figure, dict):
                    raise JobError(f"{prefix}.figures[{figure_index}] must be an object.")
                text(figure.get("path"), f"{prefix}.figures[{figure_index}].path")
                text(figure.get("caption"), f"{prefix}.figures[{figure_index}].caption")
    return sections


def add_steps(doc: Document, job_path: Path, sections: list[dict[str, Any]]) -> int:
    step_number = 0
    for section_index, section in enumerate(sections, start=1):
        add_heading(doc, f"{section_index}. {text(section['title'], 'section.title')}")
        for step in section["steps"]:
            step_number += 1
            paragraph = doc.add_paragraph()
            set_paragraph(paragraph, after=2, line=1.25)
            add_run(paragraph, f"{step_number}. {text(step['title'], 'step.title')}：", size=10.8, color=INK, bold=True)
            add_run(paragraph, text(step["action"], "step.action"), size=10.8)
            result = doc.add_paragraph()
            result.paragraph_format.left_indent = Inches(0.25)
            set_paragraph(result, after=5, line=1.2)
            add_run(result, "预期结果：", size=9.7, color=DARK_BLUE, bold=True)
            add_run(result, text(step["expected_result"], "step.expected_result"), size=9.7, color=MUTED)
            for note in text_array(step.get("notes"), "step.notes"):
                add_note(doc, "注意事项", note)
            for figure in step.get("figures", []):
                add_figure(doc, job_path, figure)
    return step_number


def add_issues(doc: Document, job: dict[str, Any]) -> None:
    issues = job.get("common_issues", [])
    if not issues:
        return
    if not isinstance(issues, list):
        raise JobError("common_issues must be an array.")
    add_heading(doc, "常见问题", 1)
    for index, issue in enumerate(issues):
        if not isinstance(issue, dict):
            raise JobError(f"common_issues[{index}] must be an object.")
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, after=5, line=1.2)
        add_run(paragraph, f"{text(issue.get('issue'), f'common_issues[{index}].issue')}：", size=10.3, color=DARK_BLUE, bold=True)
        add_run(paragraph, text(issue.get("resolution"), f"common_issues[{index}].resolution"), size=10.3)


def add_versions(doc: Document, job: dict[str, Any]) -> None:
    versions = job.get("version_history", [{"version": "1.0", "date": "待确认", "change": "首次发布。"}])
    if not isinstance(versions, list) or not versions:
        raise JobError("version_history must be a non-empty array when provided.")
    add_heading(doc, "版本记录", 1)
    table = doc.add_table(rows=1, cols=3)
    prepare_table(table, [1500, 2200, 5660])
    headers = ("版本", "日期", "变更说明")
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, HEADER_FILL)
        paragraph = clear_cell(cell)
        set_paragraph(paragraph, after=0, line=1.15)
        add_run(paragraph, value, size=9.4, color=DARK_BLUE, bold=True)
    mark_header(table.rows[0])
    for index, item in enumerate(versions):
        if not isinstance(item, dict):
            raise JobError(f"version_history[{index}] must be an object.")
        row = table.add_row()
        prepare_table(table, [1500, 2200, 5660])
        for cell, value in zip(
            row.cells,
            (
                text(item.get("version"), f"version_history[{index}].version"),
                text(item.get("date"), f"version_history[{index}].date"),
                text(item.get("change"), f"version_history[{index}].change"),
            ),
        ):
            paragraph = clear_cell(cell)
            set_paragraph(paragraph, after=0, line=1.15)
            add_run(paragraph, value, size=9.2)


def build(job_path: Path) -> Path:
    job = load_job(job_path)
    reject_credentials(job)
    for field in ("title", "system_name", "applicable_scope", "applicable_users"):
        text(job.get(field), field)
    output_path = resolve_path(job_path, job.get("output_path"), "output_path")
    if output_path.suffix.lower() != ".docx":
        raise JobError("output_path must end with .docx.")
    sections = validate_sections(job)
    preparation = text_array(job.get("preparation"), "preparation")
    notes = text_array(job.get("notes"), "notes")
    doc = Document()
    configure_document(doc, text(job["title"], "title"))
    add_cover(doc, job)
    add_heading(doc, "1. 目的与适用范围")
    add_body(doc, f"本指南用于{ text(job['applicable_scope'], 'applicable_scope') }，适用于{ text(job['applicable_users'], 'applicable_users') }。")
    add_heading(doc, "2. 操作前准备")
    for item in preparation or ["确认已进入目标环境，并具备完成本流程所需权限。"]:
        add_bullet(doc, item)
    for note in notes:
        add_note(doc, "注意事项", note)
    add_heading(doc, "3. 操作步骤")
    add_steps(doc, job_path, sections)
    add_issues(doc, job)
    add_versions(doc, job)
    doc.core_properties.title = text(job["title"], "title")
    doc.core_properties.subject = "网页后台标准操作规程"
    doc.core_properties.author = "Operations Team"
    doc.core_properties.comments = "Generated from confirmed screenshots; no credential values are stored."
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main() -> int:
    args = parse_args()
    try:
        output_path = build(Path(args.job).resolve())
        print(json.dumps({"status": "ok", "output": str(output_path)}, ensure_ascii=False))
        return 0
    except JobError as error:
        print(f"DOCX build failed: {error}")
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
